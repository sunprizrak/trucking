from docx import Document

from trucking.settings import MEDIA_ROOT

document = Document(f'{MEDIA_ROOT}/Заявка_№5_Wuxi.docx')

paragraphs = document.paragraphs
table = document.tables

print(table[0].rows[1].cells[2].text)
